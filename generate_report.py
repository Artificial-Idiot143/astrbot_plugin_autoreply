import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           "群聊自动回复系统研究报告.docx")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for attr in ['sz', 'val', 'color', 'space']:
                if attr in edge_data:
                    element.set(qn(f'w:{attr}'), str(edge_data[attr]))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(13)
    return heading


def add_body_para(doc, text, bold=False, indent=True):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(6)
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    run = para.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.bold = bold
    return para


def add_code_block(doc, code_text):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.2
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.left_indent = Cm(1)
    run = para.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return para


def build():
    doc = Document()

    # ========== 页面设置 ==========
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ========== 封面 ==========
    for _ in range(4):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('本科生\u201c科学与社会\u201d研讨课\n研究报告')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(26)
    run.bold = True

    doc.add_paragraph()

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("——基于大语言模型的群聊自动回复系统\n以记忆库为核心的设计与实现")
    run.font.name = '楷体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.size = Pt(18)

    for _ in range(3):
        doc.add_paragraph()

    info_items = [
        "学    院：________________________",
        "学生姓名：________________________",
        "学    号：________________________",
        "指导教师：________________________",
        "日    期：________________________",
    ]
    for item in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(14)

    doc.add_page_break()

    # ========== 一、研究工作计划 ==========
    add_heading_styled(doc, "一、研究工作计划", level=1)

    table = doc.add_table(rows=6, cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["时间节点", "工作内容", "阶段目标", "完成情况"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(11)
                run.bold = True

    plan_data = [
        ["第1-2周\n（选题与调研）",
         "调研群聊机器人现状；\n确定技术方案与架构",
         "完成需求分析文档；\n确定LLM选型",
         "已完成"],
        ["第3-6周\n（核心开发）",
         "实现记忆库存储引擎；\n开发消息清洗与关键词提取；\n搭建自动回复流水线",
         "记忆库CRUD完成；\n回复引擎可运行",
         "已完成"],
        ["第7-10周\n（功能完善）",
         "实现三层分级记忆管理；\n开发日常维护与趋势分析；\n思维链剥离与文本后处理",
         "记忆维护自动化；\n回复质量达标",
         "已完成"],
        ["第11-14周\n（测试优化）",
         "编写审计测试套件；\n修复Bug与性能调优；\n撰写研究报告",
         "100项测试全部通过；\n报告完成",
         "已完成"],
        ["第15-16周\n（答辩准备）",
         "准备答辩PPT；\n模拟答辩演练",
         "答辩材料就绪",
         "进行中"],
    ]
    for r, row_data in enumerate(plan_data):
        for c, text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10)

    doc.add_page_break()

    # ========== 二、摘要与关键词 ==========
    add_heading_styled(doc, "二、摘要与关键词", level=1)

    add_body_para(doc, "摘要：", bold=True, indent=False)
    abstract = (
        "本研究设计并实现了一个基于大语言模型（LLM）的QQ群聊自动回复系统，"
        "其核心是一个具备三层分级管理、智能衰减评分和检索驱动剪枝（QDP）机制的"
        "结构化记忆库。系统通过消息清洗、关键词提取、事件抽取等NLP流水线将群聊"
        "内容转化为结构化记忆，并利用思维链剥离与候选选择算法确保回复的简洁性与"
        "准确性。记忆库采用Tier 0至Tier 3的四级时间分层架构，结合Intelligent Decay"
        "与FadeMem多因素评分算法，实现了记忆的自动迭代清理与重要性排序。系统还"
        "集成了日常维护模块，支持记忆平移、时段合并与五维度趋势分析（热词检测、"
        "新词发现、句式风格变化、话题热度迁移、情绪周期）。经100项审计测试验证，"
        "系统在回复质量、记忆检索准确率和运行稳定性方面均达到设计预期，为群聊场景"
        "下的智能交互提供了一套完整的工程解决方案。"
    )
    add_body_para(doc, abstract)

    add_body_para(doc, "关键词：", bold=True, indent=False)
    add_body_para(doc, "大语言模型；自动回复系统；记忆库；智能衰减算法；群聊机器人")

    doc.add_page_break()

    # ========== 三、导师评审意见（留空） ==========
    add_heading_styled(doc, "三、导师评审意见", level=1)
    add_body_para(doc, "（由导师填写）", indent=False)
    for _ in range(6):
        doc.add_paragraph()
    add_body_para(doc, "导师签字：________________    日期：________________", indent=False)

    doc.add_page_break()

    # ========== 四、研究报告正文 ==========
    add_heading_styled(doc, "四、研究报告正文", level=1)

    # ---- 1. 背景和目标 ----
    add_heading_styled(doc, "1  背景和目标", level=2)

    add_heading_styled(doc, "1.1  研究背景", level=3)
    add_body_para(doc,
        "随着即时通讯工具在日常生活中的普及，QQ群聊已成为人们社交、学习与工作协作的重要平台。"
        "然而，群聊中信息量巨大、话题切换频繁，用户常常面临『消息过载』的困境——错过重要讨论、"
        "难以追踪话题脉络、无法及时回应他人消息。传统的群聊机器人多基于关键词匹配或固定规则，"
        "缺乏对上下文的理解能力，回复生硬且无法适应多样化的对话场景。"
    )
    add_body_para(doc,
        "近年来，以GPT、Qwen等为代表的大语言模型（Large Language Model, LLM）在自然语言理解"
        "与生成方面取得了突破性进展，使得构建具备上下文感知能力的智能回复系统成为可能。然而，"
        "直接将LLM应用于群聊场景面临两大核心挑战：其一，LLM的上下文窗口有限，无法容纳长时间的"
        "群聊历史；其二，LLM的输出可能包含思维链、英文混杂、标点粘连等问题，需要有效的后处理"
        "机制来保证回复质量。"
    )

    add_heading_styled(doc, "1.2  研究目标", level=3)
    add_body_para(doc,
        "本研究旨在设计并实现一个以结构化记忆库为核心的群聊自动回复系统，具体目标包括："
    )
    goals = [
        "构建一个具备消息存储、关键词索引、事件抽取能力的记忆库引擎，支持高效检索与自动维护；",
        "设计三层分级记忆管理策略，结合智能衰减算法实现记忆的自动迭代清理，防止数据库无限膨胀；",
        "开发完整的自动回复流水线，包括消息清洗、回复决策、记忆检索、LLM调用与思维链后处理；",
        "实现日常维护与趋势分析模块，支持记忆平移、时段合并及多维度群聊趋势洞察；",
        "建立完善的测试体系，确保系统各模块的正确性与稳定性。",
    ]
    for g in goals:
        add_body_para(doc, "（" + str(goals.index(g) + 1) + "）" + g)

    # ---- 2. 研究报告正文 ----
    add_heading_styled(doc, "2  系统设计与实现", level=2)

    add_heading_styled(doc, "2.1  系统总体架构", level=3)
    add_body_para(doc,
        "系统采用模块化架构设计，由五个核心模块组成：配置管理模块（config.py）、记忆库引擎"
        "（memory_ai.py）、自动回复引擎（auto_reply.py）、日常维护模块（daily_maintenance.py）"
        "以及趋势分析模块（trend_analyzer.py）。各模块之间通过明确的接口进行协作，形成一条"
        "从消息接收到回复生成的完整处理流水线。"
    )
    add_body_para(doc,
        "系统的工作流程如下：当用户消息到达时，自动回复引擎首先进行消息清洗与回复必要性判定；"
        "若需要回复，则根据消息类型决定是否检索记忆库中的历史上下文；随后调用本地部署的LLM"
        "（Qwen 3.5-9B）生成候选回复；最后通过思维链剥离与文本后处理得到最终回复。与此同时，"
        "每条消息都会异步存入记忆库，经过关键词提取与事件抽取后形成结构化记忆。日常维护模块"
        "定期执行记忆平移与清理，确保记忆库的健康运行。"
    )

    add_heading_styled(doc, "2.2  记忆库核心设计", level=3)
    add_body_para(doc,
        "记忆库是本系统的核心组件，基于SQLite实现，采用WAL（Write-Ahead Logging）模式以支持"
        "高并发读写。数据库包含五张核心表：messages（消息存储）、keywords（关键词索引）、"
        "events（事件记录）、retrieval_log（检索日志）和patterns（行为规律）。"
    )

    add_heading_styled(doc, "2.2.1  三层分级存储架构", level=3)
    add_body_para(doc,
        "记忆库采用Tier 0至Tier 3的四级时间分层架构，根据消息的时间远近采用不同的存储粒度："
    )
    tiers = [
        "Tier 0（当天）：保留原始消息全文，支持逐条检索，是回复生成的主要信息来源；",
        "Tier 1（1-2天）：按30分钟窗口合并消息，保留压缩摘要与关键词，减少存储开销；",
        "Tier 2（3-6天）：按12小时窗口合并，进一步压缩，仅保留高度概括的摘要；",
        "Tier 3（7天以上）：按天合并，仅保留事件索引与关键词，原文不再保留。",
    ]
    for t in tiers:
        add_body_para(doc, t)

    add_body_para(doc,
        "这种分层设计模拟了人类记忆的『遗忘曲线』——近期记忆细节丰富，远期记忆逐渐模糊为"
        "抽象概括，在保证信息可用性的同时有效控制了存储成本。"
    )

    add_heading_styled(doc, "2.2.2  智能衰减与重要性评分", level=3)
    add_body_para(doc,
        "系统实现了两种互补的记忆评分算法：Intelligent Decay（智能衰减）和FadeMem（多因素"
        "重要性评分）。Intelligent Decay算法基于艾宾浩斯遗忘曲线，对记忆进行时间衰减加权——"
        "越久远的记忆权重越低。FadeMem算法则综合考虑消息长度、关键词密度、用户互动频率、"
        "检索命中次数等多个维度，为每条记忆计算综合重要性分数。"
    )
    add_body_para(doc,
        "在记忆清理阶段，系统根据综合评分对低价值记忆进行淘汰，同时采用QDP（Query-Driven "
        "Pruning，检索驱动剪枝）策略——被频繁检索的记忆即使时间久远也会获得额外加分，避免"
        "重要历史信息被误删。"
    )

    add_heading_styled(doc, "2.2.3  关键词提取与事件抽取", level=3)
    add_body_para(doc,
        "消息存入记忆库时，系统调用LLM进行两项NLP任务：关键词提取和事件抽取。关键词提取"
        "从消息中识别核心概念词，建立倒排索引以支持高效检索。事件抽取则将消息归类到具体的"
        "事件主题（subject）下，并生成一句话摘要（summary），形成结构化的『事件-时间-人物』"
        "三元组。这种结构化表示使得记忆检索不仅能按关键词匹配，还能按事件脉络进行语义关联。"
    )

    add_heading_styled(doc, "2.3  自动回复引擎", level=3)

    add_heading_styled(doc, "2.3.1  消息清洗与回复决策", level=3)
    add_body_para(doc,
        "回复引擎采用四步决策流水线：第一步（step1_should_reply）判断消息是否需要回复，"
        "包括检测空消息、无意义词（如『嗯嗯』）、艾特别人（非@机器人）等不应回复的场景；"
        "第二步（step2_need_memory）判断是否需要检索历史记忆，对于简单寒暄（如『你好』）"
        "直接回复，对于概括性问题（如『聊了些什么』）或时间相关问题（如『昨天发生了什么』）"
        "则触发记忆检索。"
    )

    add_heading_styled(doc, "2.3.2  思维链剥离与文本后处理", level=3)
    add_body_para(doc,
        "由于本系统使用的Qwen 3.5-9B为思维链（Chain-of-Thought）模型，其输出中常包含"
        "英文思考过程、系统提示词回显等噪声。为此，系统实现了专门的思维链剥离函数"
        "（_strip_thinking_process），通过识别思维链结构标记（如Thinking Process、"
        "Role:、Task:等）和中文占比分析，将思考过程与最终答案分离。剥离后的文本"
        "再经过_super_clean函数进行深度清理，包括：移除残留英文、修复标点粘连（如'.，'→'.'）、"
        "统一全角标点、过滤非白名单字符等。"
    )

    add_heading_styled(doc, "2.3.3  候选选择策略", level=3)
    add_body_para(doc,
        "思维链剥离后可能产生多个候选文本片段。系统采用三阶段候选选择策略：优先选择最后一个"
        "包含足够中文字符的候选（因为最终答案通常在思考过程之后）；若无满足条件的候选，则选择"
        "中文字符最多的候选；最后回退到第一个候选。这种策略有效避免了将思考过程误认为最终回复。"
    )

    add_heading_styled(doc, "2.4  日常维护与趋势分析", level=3)
    add_body_para(doc,
        "日常维护模块（daily_maintenance.py）负责记忆库的定期健康管理，主要执行以下操作："
        "（1）Tier 0语言分析——对当天消息进行语言特征分析，提取群聊风格描述供回复引擎参考；"
        "（2）记忆平移——将消息按时间从低Tier向高Tier迁移，逐步压缩存储粒度；"
        "（3）时段合并——对同一时间窗口内的多条消息调用LLM进行智能合并，生成统一摘要；"
        "（4）数据库优化——重建索引、清理碎片、更新统计信息。"
    )
    add_body_para(doc,
        "趋势分析模块（trend_analyzer.py）提供五个维度的群聊趋势洞察：热词检测（对比近期与"
        "基线期的词频变化）、新词发现（检测基线期未出现的新兴词汇）、句式风格变化（分析平均"
        "句长、感叹号/问号占比的时序变化）、话题热度迁移（追踪各话题的讨论热度升降趋势）以及"
        "情绪周期（按小时和星期统计群聊情绪的正负峰值）。分析结果以JSON格式存入trend_log表，"
        "支持历史回溯。"
    )

    add_heading_styled(doc, "2.5  实验与测试", level=3)
    add_body_para(doc,
        "为验证系统的正确性与稳定性，本研究编写了全面的审计测试套件（audit_test.py），"
        "覆盖14个测试类别共100项测试用例，包括："
    )
    tests = [
        "文本清理测试（_super_clean）：验证英文移除、标点粘连修复、白名单过滤等功能；",
        "回复决策测试（step1_should_reply）：覆盖空消息、无意义词、@提及、艾特别人等场景；",
        "记忆需求判定测试（step2_need_memory）：验证寒暄、疑问、概括性问题的分类准确性；",
        "记忆简单总结测试（_simple_summary_from_memory）：验证多记忆拼接与英文清理；",
        "智能默认回复测试（_smart_default_reply）：验证无记忆时的回退策略；",
        "时间关键词解析测试（parse_time_keywords）：验证『今天』『上周』等时间表达；『』"
        "思维链剥离测试（_strip_thinking_process）：验证Final Answer提取与候选选择；",
        "二值判定提取测试（extract_binary_decision）：验证LLM输出的0/1判定解析；",
        "数据库表结构检查：验证所有表与索引的存在性和列完整性；",
        "维护流程一致性检查：验证记忆平移各函数的可调用性；",
        "死代码检查：确认已废弃函数已被移除；",
        "边界条件测试：超长字符串、纯标点、混合中英文标点等极端输入；",
        "端到端模拟测试：10种典型对话场景的完整流水线验证。",
    ]
    for t in tests:
        add_body_para(doc, "（" + str(tests.index(t) + 1) + "）" + t)

    add_body_para(doc,
        "测试结果显示，全部100项测试用例均通过，通过率100%。系统在消息清洗、回复决策、"
        "记忆检索、思维链剥离、数据库完整性等关键路径上均表现正确。"
    )

    # ---- 3. 结论/总结 ----
    add_heading_styled(doc, "3  结论与展望", level=2)

    add_heading_styled(doc, "3.1  工作总结", level=3)
    add_body_para(doc,
        "本研究成功设计并实现了一个以结构化记忆库为核心的群聊自动回复系统。主要贡献包括："
    )
    contributions = [
        "提出并实现了三层分级记忆存储架构，模拟人类遗忘曲线，在信息保真度与存储效率之间取得平衡；",
        "设计了Intelligent Decay与FadeMem双算法评分机制，结合QDP检索加权策略，实现了智能化的记忆迭代清理；",
        "构建了完整的自动回复流水线，包括消息清洗、回复决策、LLM调用与思维链后处理，确保回复的准确性与简洁性；",
        "开发了日常维护与五维度趋势分析模块，使系统具备自我健康管理能力；",
        "建立了100项审计测试的验证体系，覆盖所有核心模块与边界条件。",
    ]
    for c in contributions:
        add_body_para(doc, "（" + str(contributions.index(c) + 1) + "）" + c)

    add_heading_styled(doc, "3.2  未来展望", level=3)
    add_body_para(doc,
        "本系统仍存在若干可改进方向：（1）当前记忆检索主要依赖关键词匹配，可引入向量嵌入"
        "（Embedding）实现语义级相似度检索，提升长尾问题的召回率；（2）回复生成目前为单轮"
        "调用，可引入多轮反思机制让LLM自我修正回复；（3）趋势分析模块的情感词典较为简单，"
        "可替换为专业情感分析模型以提高准确性；（4）系统目前依赖本地LLM部署，可适配云端"
        "API以降低硬件门槛。"
    )

    # ---- 4. 致谢 ----
    add_heading_styled(doc, "4  致谢", level=2)
    add_body_para(doc,
        "本研究的完成离不开指导教师的悉心指导和宝贵建议。在系统设计与实现过程中，"
        "指导教师对技术方案的可行性、系统架构的合理性以及实验验证的严谨性均给予了"
        "关键性的指导。同时，感谢课程提供的学习平台与实践机会，使本人能够将课堂所学的"
        "编译原理、自然语言处理等理论知识应用于实际工程问题。此外，感谢开源社区提供的"
        "Qwen大语言模型、SQLite数据库等优秀工具，为本研究的顺利开展提供了坚实的技术基础。"
    )

    # ---- 5. 附录 ----
    add_heading_styled(doc, "5  附录", level=2)

    add_heading_styled(doc, "附录A  项目文件结构", level=3)
    add_body_para(doc, "项目主要源文件及其功能说明：", indent=False)

    files = [
        ("config.py", "系统配置文件，包含数据库路径、LLM API地址与模型名称"),
        ("memory_ai.py", "记忆库引擎核心，实现MemoryDB类、消息存储、关键词提取、事件抽取、记忆清理与合并"),
        ("auto_reply.py", "自动回复引擎，实现消息清洗、回复决策、LLM调用、思维链剥离与文本后处理"),
        ("daily_maintenance.py", "日常维护模块，实现记忆平移（Tier迁移）、时段合并、语言分析与数据库优化"),
        ("trend_analyzer.py", "趋势分析模块，实现热词检测、新词发现、句式风格变化、话题热度迁移、情绪周期五维度分析"),
        ("reply_engine.py", "独立回复引擎（含推测引擎与规律匹配），提供另一种回复生成策略"),
        ("batch_import.py", "历史记录批量导入脚本，支持断点续跑与分级存储"),
        ("batch_import1.py", "Tier3专用快速导入脚本，采用批量LLM调用优化性能"),
        ("audit_test.py", "审计测试套件，包含14类100项自动化测试"),
        ("test_reply.py", "交互式回复测试工具，支持终端对话调试"),
    ]
    for fname, fdesc in files:
        add_body_para(doc, f"· {fname}：{fdesc}")

    add_heading_styled(doc, "附录B  数据库表结构", level=3)
    add_body_para(doc, "记忆库（memory.db）包含以下核心数据表：", indent=False)

    tables = [
        ("messages", "消息主表，字段：id, user_id, user_name, content, compressed_content, content_state, importance_score, tier, last_maintained, timestamp"),
        ("keywords", "关键词索引表，字段：id, msg_id, keyword"),
        ("events", "事件记录表，字段：id, msg_id, subject, summary, user_name, timestamp"),
        ("retrieval_log", "检索日志表，字段：id, msg_ids, query_text, timestamp"),
        ("patterns", "行为规律表，字段：id, trigger_words, pattern_desc, confidence, last_updated"),
        ("trend_log", "趋势分析日志表，字段：id, analysis_time, period_start, period_end, analysis_json"),
    ]
    for tname, tdesc in tables:
        add_body_para(doc, f"· {tname}：{tdesc}")

    # ---- 6. 参考文献 ----
    add_heading_styled(doc, "6  参考文献", level=2)

    refs = [
        "[1] Vaswani A, Shazeer N, Parmar N, et al. Attention is All You Need[C]. "
        "Advances in Neural Information Processing Systems, 2017: 5998-6008.",

        "[2] Brown T B, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[C]. "
        "Advances in Neural Information Processing Systems, 2020: 1877-1901.",

        "[3] Wei J, Wang X, Schuurmans D, et al. Chain-of-Thought Prompting Elicits "
        "Reasoning in Large Language Models[C]. Advances in Neural Information Processing "
        "Systems, 2022: 24824-24837.",

        "[4] Ebbinghaus H. Memory: A Contribution to Experimental Psychology[M]. "
        "New York: Teachers College Press, 1913.",

        "[5] Anderson J R, Schooler L J. Reflections of the Environment in Memory[J]. "
        "Psychological Science, 1991, 2(6): 396-408.",

        "[6] Qwen Team. Qwen Technical Report[EB/OL]. https://github.com/QwenLM/Qwen, 2024.",

        "[7] SQLite Consortium. SQLite Documentation[EB/OL]. "
        "https://www.sqlite.org/docs.html, 2024.",

        "[8] 张昱, 胡燕, 刘鹏. 编译原理课程设计[M]. 北京: 机械工业出版社, 2019.",
    ]
    for ref in refs:
        add_body_para(doc, ref)

    # ========== 保存 ==========
    doc.save(OUTPUT_PATH)
    print(f"报告已生成: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()